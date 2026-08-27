#!/usr/bin/env python3
"""Extract text from PDF/DOCX files in source-pdfs/ into structured
Markdown in wiki-pages/.

For each source document:
  1. PDF: try the text layer with pdfplumber, per page — any page whose own
     text is missing or broken (too little text, or mostly garbage
     characters) gets OCR'd individually (pdf2image -> tesseract, ukr+eng),
     while pages with a good text layer keep it. This matters because a
     mixed PDF (a scanned cover page glued onto an otherwise-native
     document) would otherwise pass a whole-document average check and
     silently lose just that one page. DOCX: read directly with
     python-docx — it always has a real text layer, no OCR needed.
  2. Split the text into logical blocks by document structure (roman-numeral
     sections, numbered "punkty"/"pidpunkty", etc.) using Markdown headings.
  3. Try to recognize order number / date from the text.
  4. Write <stem>.md into wiki-pages/ with a YAML frontmatter header.

Usage:
    .venv/bin/python scripts/extract_to_markdown.py
"""

from __future__ import annotations

import logging
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from urllib.parse import quote

import docx
import pdfplumber
import pytesseract
import yaml
from pdf2image import convert_from_path

ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "source-pdfs"
OUTPUT_DIR = ROOT / "wiki-pages"

# Cross-document links are written as plain Markdown (not Liquid), so they
# can't use Jekyll's `relative_url` filter — this must match `baseurl` in
# _config.yml, since the site is served from a /docs-rag-pilot/ subpath
# (a GitHub Pages project site), not the domain root.
SITE_BASEURL = "/docs-rag-pilot"

OCR_LANGS = "ukr+eng"
# 400 (up from 300): the cover-page seal/stamp overlaps the ЗАТВЕРДЖЕНО block
# on some scans, and at 300 DPI tesseract garbles text right at that overlap
# (the anchor word, the issuing-naказ line) even though the title's actual
# subject a line or two below comes through fine either way. The higher
# resolution recovers several of those anchor-word lines too.
OCR_DPI = 400

# Below this many characters we treat one page's own text layer as
# missing/broken and OCR that page individually.
MIN_CHARS_PER_PAGE = 40

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger("extract_to_markdown")


@dataclass
class ExtractionResult:
    source_path: Path
    text: str
    pages_text: list[str]
    used_ocr: bool
    pages: int
    error: str | None = None


@dataclass
class RunStats:
    total: int = 0
    succeeded: int = 0
    used_ocr: int = 0
    failed: list[tuple[str, str]] = field(default_factory=list)


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------

# A raw text line from pdfplumber/OCR is a physical line wrapped at the
# page's text-box width, not a logical paragraph — so consecutive sentences
# that were separate indented paragraphs in the original document end up as
# consecutive lines with no blank line between them, and downstream
# Markdown rendering merges them into one wall of text. The fix is the same
# one a human proofreader would use: a paragraph's *first* line starts
# further right (the "абзацний відступ" indent) than every line after it.
# Below this many points of difference from a page's baseline left margin,
# treat it as page-to-page jitter rather than a real indent.
INDENT_TOLERANCE_PT = 3.0

# Not every document in this corpus signals a new paragraph with a visual
# first-line indent — some number every point ("1.1. ...", "1.2. ...") and
# start each one flush with the same left margin as its own continuation
# lines, relying on the number itself (not indentation) as the paragraph
# marker. A bullet character works the same way. Recognize both as a
# paragraph start regardless of x0, alongside the indent signal -- this
# mirrors RE_SECTION/RE_POINT further below, which recognize the same
# numbering as *headings*; here it is the coarser paragraph-vs-continuation
# split. U+F0B7 is Wingdings/Symbol-font bullet glyph, which Word
# documents commonly embed as a private-use-area codepoint instead of a
# plain bullet character.
RE_PARAGRAPH_MARKER = re.compile(
    r'^(?:[IVXLCDM]{1,6}\.\s|\d{1,3}(?:\.\d{1,3}){0,3}\.?\s|\d{1,2}\)\s|[а-яіїєґ]\)\s|[-−–—•·]\s)'
)


def _reflow_lines(lines: list[tuple[float, str]]) -> str:
    """lines: (x0, text) pairs in reading order. Joins each paragraph's
    wrapped continuation lines onto one line with a space, and separates
    distinct paragraphs (and anything else that starts at a different
    indent, or with its own numbered/bulleted marker — headings, list
    items, the repeated page header/footer) with a blank line, so each
    becomes its own Markdown paragraph/block instead of all merging into
    one."""
    lines = [(x0, text.strip()) for x0, text in lines if text and text.strip()]
    if not lines:
        return ""

    # The baseline (continuation) margin is the leftmost x0 that recurs —
    # picking the single most *frequent* x0 instead can pick the paragraph
    # indent by mistake on a page made up mostly of short, one-line
    # paragraphs, where the indent position outnumbers the baseline one.
    from collections import Counter
    counts = Counter(round(x0) for x0, _ in lines)
    recurring = [x for x, n in counts.items() if n >= 2]
    baseline = min(recurring) if recurring else round(lines[0][0])

    paragraphs: list[str] = []
    buffer = ""
    for x0, text in lines:
        is_new_paragraph = (
            abs(round(x0) - baseline) > INDENT_TOLERANCE_PT
            or bool(RE_PARAGRAPH_MARKER.match(text))
        )
        if buffer and not is_new_paragraph:
            buffer += " " + text
        else:
            if buffer:
                paragraphs.append(buffer)
            buffer = text
    if buffer:
        paragraphs.append(buffer)

    return "\n\n".join(paragraphs)


def extract_with_pdfplumber(pdf_path: Path) -> list[tuple[str, str]]:
    """Returns (raw, reflowed) text for each page. Raw is pdfplumber's
    plain per-physical-line extract_text() — what guess_title/
    guess_order_number/guess_order_date/build_link_spec all expect, since
    they treat a blank line as the firm end of a title/approval block, and
    reflow would insert one between every wrapped line the block happens
    to be made of. Reflowed collapses each paragraph's wrapped lines onto
    one line and separates paragraphs with a blank line instead — for the
    body, where a reader wants actual paragraphs, not the PDF's line
    wrapping."""
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            raw = page.extract_text() or ""
            try:
                lines = page.extract_text_lines()
            except Exception:
                lines = None
            reflowed = _reflow_lines([(l["x0"], l["text"]) for l in lines]) if lines else raw
            pages.append((raw, reflowed))
    return pages


def _reflow_ocr_image(image) -> str:
    """Tesseract's own layout analysis already groups words into
    paragraphs (block_num/par_num) — use that instead of inventing an x0
    heuristic a second time for OCR'd pages."""
    from pytesseract import Output

    data = pytesseract.image_to_data(image, lang=OCR_LANGS, output_type=Output.DICT)
    paragraphs: dict[tuple[int, int, int], list[str]] = {}
    order: list[tuple[int, int, int]] = []
    for i, word in enumerate(data["text"]):
        if not word.strip():
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        # Group by (block, par) for paragraphs, but preserve line order
        # within a paragraph via line_num when joining below.
        para_key = (data["block_num"][i], data["par_num"][i])
        if para_key not in paragraphs:
            paragraphs[para_key] = []
            order.append(para_key)
        paragraphs[para_key].append((data["line_num"][i], word))

    out = []
    for key in order:
        words = paragraphs[key]
        # Reconstruct with line breaks collapsed into spaces, preserving
        # the order Tesseract emitted the words in (already line-major).
        out.append(" ".join(w for _, w in words))
    text = "\n\n".join(out)
    return text if text.strip() else pytesseract.image_to_string(image, lang=OCR_LANGS)


def _ocr_page_text(image) -> tuple[str, str]:
    """Returns (raw, reflowed) OCR text for one page image — same raw-vs-
    reflowed split as extract_with_pdfplumber, for the same reason."""
    raw = pytesseract.image_to_string(image, lang=OCR_LANGS)
    reflowed = _reflow_ocr_image(image)
    return raw, reflowed


def extract_with_ocr(pdf_path: Path) -> list[tuple[str, str]]:
    """Whole-document OCR — only used when pdfplumber can't open the PDF at
    all (malformed file), so there's no per-page text to decide from."""
    images = convert_from_path(str(pdf_path), dpi=OCR_DPI)
    return [_ocr_page_text(img) for img in images]


def ocr_single_page(pdf_path: Path, page_number: int) -> tuple[str, str]:
    """page_number is 1-indexed. Rasters just this one page instead of the
    whole document — cheap even for a 30+ page PDF where only the cover
    needs OCR."""
    images = convert_from_path(
        str(pdf_path), dpi=OCR_DPI, first_page=page_number, last_page=page_number
    )
    return _ocr_page_text(images[0])


def extract_with_docx(docx_path: Path) -> tuple[str, int]:
    """.docx always carries a real text layer, so no OCR fallback is ever
    needed. Table cells are pulled in too (flattened, cell by cell) since
    some normative content lives in tables rather than paragraphs."""
    document = docx.Document(docx_path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts), len(document.paragraphs)


def is_page_broken(text: str) -> bool:
    """Per-page version of the old whole-document check. A mixed PDF (most
    pages a real text layer, one page a scanned cover) used to slip past
    detection entirely: the broken page's near-empty text got averaged in
    with 20+ healthy pages, so the document-level average never dropped
    below MIN_CHARS_PER_PAGE and that one page was silently left blank."""
    stripped = (text or "").strip()
    if not stripped:
        return True
    if len(stripped) < MIN_CHARS_PER_PAGE:
        return True
    # Mostly non-letter characters usually means a garbled/font-mapped layer.
    letters = sum(ch.isalpha() for ch in stripped)
    if letters / max(len(stripped), 1) < 0.3:
        return True
    return False


def extract_document(source_path: Path) -> ExtractionResult:
    if source_path.suffix.lower() == ".docx":
        text, n_units = extract_with_docx(source_path)
        return ExtractionResult(source_path, text, pages_text=[text], used_ocr=False, pages=n_units)

    try:
        pages = extract_with_pdfplumber(source_path)  # list[(raw, reflowed)]
    except Exception as exc:  # pdfplumber can choke on malformed PDFs
        log.warning("pdfplumber failed for %s (%s), falling back to whole-document OCR", source_path.name, exc)
        pages = extract_with_ocr(source_path)
        raw_text = "\n\n".join(raw for raw, _ in pages)
        reflowed_pages = [reflowed for _, reflowed in pages]
        return ExtractionResult(source_path, raw_text, pages_text=reflowed_pages, used_ocr=True, pages=len(pages))

    broken_pages = [i for i, (raw, _) in enumerate(pages) if is_page_broken(raw)]
    if broken_pages:
        log.info(
            "%s: %d/%d page(s) missing/weak text layer, OCR'ing just those (ukr+eng)...",
            source_path.name, len(broken_pages), len(pages),
        )
        for i in broken_pages:
            pages[i] = ocr_single_page(source_path, i + 1)

    raw_text = "\n\n".join(raw for raw, _ in pages)
    reflowed_pages = [reflowed for _, reflowed in pages]
    return ExtractionResult(source_path, raw_text, pages_text=reflowed_pages, used_ocr=bool(broken_pages), pages=len(pages))


# --------------------------------------------------------------------------
# Structure detection / Markdown formatting
# --------------------------------------------------------------------------

# This university's documents share one ISO-9001 template: every single
# page repeats an institutional header/footer stamp, and the cover page
# carries an approval block (ЗАТВЕРДЖЕНО...) whose useful parts (title,
# наказ number/date) are already pulled into the YAML frontmatter by
# guess_title/guess_order_number/guess_order_date. Both are pure noise in
# the document body and get stripped by strip_boilerplate() below.
RUNNING_HEADER_PATTERNS = [
    re.compile(r"МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ"),
    re.compile(r"ДЕРЖАВНИЙ УНІВЕРСИТЕТ.*ЖИТОМИРСЬКА ПОЛІТЕХНІКА"),
    re.compile(r"Система управління якістю.*ДСТУ ISO"),
    re.compile(r"Екземпляр\s*№?\s*\d+.*Арк"),
    re.compile(r"^Випуск\s+\d+\s+Зміни\s+\d+"),
    # Document control code, e.g. "П-10.00-02.01-" / "07-2025" /
    # "02.15-01-2020", which different documents wrap onto their own
    # line(s) in slightly different ways instead of the header line above.
    # A standalone line made up only of digits/dots/dashes (with an
    # optional single leading Cyrillic letter) is never real prose. This
    # intentionally requires a trailing digit/dash (not a bare dot) so a
    # lone clause number like "3.2.13." is never matched here.
    re.compile(r"^(?:[А-ЯІЇЄҐ]-)?\d+(?:[.\-]\d+)*-?$"),
    # The same code, but with stray spaces OCR sometimes inserts around
    # separators ("П-10.00-02.18- 01-2024") — only fires when the leading
    # letter marker is actually present, so it can never match a bare
    # clause number (which never carries one).
    re.compile(r"^[А-ЯІЇЄҐ]-\s*[\d./\-\s]+$"),
]
# The institution's logo text sometimes extracts as isolated lines.
STANDALONE_NOISE_LINES = {"Житомирська", "політехніка"}

# Table-of-contents dot leaders ("Загальні положення……………… 3") and bare
# page-footer numbers.
RE_DOT_LEADER = re.compile(r"[.…]{3,}\s*\d*\s*$")
RE_LONE_PAGE_NUMBER = re.compile(r"^\d{1,3}$")

# The same institutional stamp above also leaks through in mixed/title
# case (title-case OCR reads, or a native text layer that stores the
# header run in title case despite it being styled as small caps on the
# page) instead of the ALL-CAPS form RUNNING_HEADER_PATTERNS expects. A
# case-insensitive version of that pattern would also eat ordinary
# sentences that just mention the university by name (very common in
# body prose), so this only fires on lines that are *nothing but* the
# institution name plus doc-code/date noise: no leading digit (real
# numbered clauses, e.g. "3.2. Міністерство освіти і науки України:",
# must never be touched) and no real word left over once the
# institution phrase and known filler words are removed.
RE_INSTITUTION_MIXED_CASE = re.compile(
    r"міністерство освіти і науки україни"
    r"|державний університет\s*[«\"“]?\s*житомирська\s*політехніка[»\"”]?",
    re.IGNORECASE,
)
RE_REAL_WORD = re.compile(r"[А-ЯІЇЄҐа-яіїєґ]{4,}")
_INSTITUTION_FILLER_WORDS = {"житомирська", "політехніка", "університет"}


def is_mixed_case_header_noise(stripped_line: str) -> bool:
    if not stripped_line or stripped_line[0].isdigit():
        return False
    if RE_INSTITUTION_MIXED_CASE.search(stripped_line):
        residual = RE_INSTITUTION_MIXED_CASE.sub(" ", stripped_line)
        leftover_words = [
            w for w in RE_REAL_WORD.findall(residual)
            if w.lower() not in _INSTITUTION_FILLER_WORDS
        ]
        return not leftover_words
    # Shorter variant of the same stamp: just "Житомирська"/"політехніка"
    # (together, or glued to a doc-code fragment like "П-14.01-01-02-2019"),
    # with no institution phrase to anchor on. Capped at 80 chars and no
    # other real word present, so a genuine sentence that happens to start
    # with these words on a wrapped PDF line is never touched.
    if len(stripped_line) > 80:
        return False
    words = RE_REAL_WORD.findall(stripped_line)
    if not words:
        return False
    return all(w.lower() in _INSTITUTION_FILLER_WORDS for w in words)

def drop_cover_page(pages_text: list[str]) -> list[str]:
    """Page 1 of a "ПОЛОЖЕННЯ/ПОЛІТИКА/..." document is a dedicated cover —
    ЗАТВЕРДЖЕНО block, title, signatures, "Контрольний примірник" — none of
    it is body content (its useful parts are already pulled into the YAML
    frontmatter by guess_title/guess_order_number/guess_order_date). Drop it
    outright rather than trying to guess where it ends line-by-line: that
    guessing is exactly what used to eat real content, like an unnumbered
    "ВСТУП" section landing on page 2/3 right after a ЗМІСТ block whose end
    was misdetected.

    A наказ, by contrast, approves *itself* — it has no separate cover, so
    page 1 (or the only page) is real content and must be kept. Whether
    page 1 is a cover is decided by whether it actually contains a
    ЗАТВЕРДЖЕНО stamp, not by position alone."""
    if len(pages_text) > 1 and RE_APPROVAL_WORD.search(pages_text[0]):
        return pages_text[1:]
    return pages_text


def strip_boilerplate(text: str) -> str:
    """Remove the repeated page header/footer stamp from raw extracted
    text. The cover page and its ЗАТВЕРДЖЕНО block are handled separately,
    by drop_cover_page(), before this ever runs — this only deals with the
    ISO-template stamp line that repeats on every remaining page."""
    out: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        if (
            stripped in STANDALONE_NOISE_LINES
            or RE_DOT_LEADER.search(stripped)
            or RE_LONE_PAGE_NUMBER.match(stripped)
            or any(p.search(stripped) for p in RUNNING_HEADER_PATTERNS)
            or is_mixed_case_header_noise(stripped)
        ):
            continue
        out.append(line)

    return "\n".join(out)


# Every document in this ISO-9001 template ends with the same tail: a
# sign-off block ("Керівник відділу МЗГ М.О. Псюк Погоджено: Перший
# проректор О.В. Олійник ...") immediately followed by a run of
# standardized tracking-sheet sections ("Аркуш поширення документа",
# "Аркуш ознайомлення з документом", "Аркуш обліку змін", "Аркуш
# реєстрації ревізій") that are blank forms, not policy content — and
# whose tabular layout regularly comes out as reflow gibberish (see the
# "и и № листа/сторінки" style garbage in ekolohichna_polityka.md). None
# of this belongs in the document body.
RE_TRAILER_SHEET = re.compile(
    r"аркуш\s+(?:поширення|ознайомлення|обліку\s+змін|реєстраці[її])",
    re.IGNORECASE,
)
# A sign-off line names people as "І.ПБ. Прізвище" (two initials, then a
# capitalized surname) — a shape that essentially never occurs in normal
# running prose (the cover page's ЗАТВЕРДЖЕНО block, the one place a name
# is written this way earlier in the document, is already dropped by
# drop_cover_page before this runs).
RE_SIGNATURE_NAME = re.compile(r"[А-ЯІЇЄҐ]\.\s?[А-ЯІЇЄҐ]\.\s*[А-ЯІЇЄҐ][а-яіїєґ'’ʼ]{2,}")


def strip_trailing_administrivia(text: str) -> str:
    """Cut everything from the first tracking-sheet header onward, plus
    the sign-off block that's typically glued onto the end of the last
    real paragraph right before it (same paragraph, no blank line, since
    the page layout puts names right after the last sentence)."""
    sheet_match = RE_TRAILER_SHEET.search(text)
    if not sheet_match:
        return text
    # Text up to the sheet header, with the run of blank lines that
    # precedes it (sometimes several, e.g. an emptied-out page) trimmed
    # off, so it ends right at the last real character.
    before = text[: sheet_match.start()].rstrip()
    # The sign-off block sits somewhere in the last stretch of `before` —
    # sometimes glued onto the last real sentence, sometimes its own
    # paragraph, sometimes followed by a further "(Ф 03.02-01)" form-code
    # paragraph before the sheet header. Rather than guess which of the
    # last few paragraphs it's in, just look for the first signature-name
    # shape within a generous trailing window.
    window_start = max(0, len(before) - 800)
    sig_match = RE_SIGNATURE_NAME.search(before, window_start)
    cutoff = sig_match.start() if sig_match else len(before)
    return text[:cutoff].rstrip()


# Roman-numeral top-level sections, e.g. "I. ЗАГАЛЬНІ ПОЛОЖЕННЯ"
RE_SECTION = re.compile(r"^([IVXLCDM]{1,6})\.\s+(\S.*)$")
# Numbered points, e.g. "1. Текст" or "1.2. Текст"
RE_POINT = re.compile(r"^(\d{1,3}(?:\.\d{1,3}){0,3})\.?\s+(\S.*)$")
# Sub-points like "1.2.3." handled by RE_POINT via dot count.

# "N. <month name>" is a date ("15 вересня 2025 р."), not a document point.
_MONTH_NAMES_LOWER = tuple(m.lower() for m in [
    "січня", "лютого", "березня", "квітня", "травня", "червня",
    "липня", "серпня", "вересня", "жовтня", "листопада", "грудня",
])


def _looks_like_date(point_body: str) -> bool:
    first_word = point_body.split(maxsplit=1)[0].lower().rstrip(",.") if point_body else ""
    return first_word in _MONTH_NAMES_LOWER


def split_into_blocks(text: str) -> str:
    """Reflow raw extracted text into Markdown headings for sections/points."""
    lines = [ln.rstrip() for ln in text.split("\n")]
    out: list[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            out.append("")
            continue

        m_section = RE_SECTION.match(line)
        if m_section:
            out.append(f"\n## {m_section.group(1)}. {m_section.group(2)}\n")
            continue

        m_point = RE_POINT.match(line)
        if m_point and not _looks_like_date(m_point.group(2)):
            depth = m_point.group(1).count(".") + 1  # 1 -> level3, 1.2 -> level4, ...
            heading_level = min(3 + (depth - 1), 6)
            out.append(f"\n{'#' * heading_level} {m_point.group(1)}. {m_point.group(2)}")
            continue

        out.append(line)

    # Collapse 3+ blank lines into a single blank line.
    result = re.sub(r"\n{3,}", "\n\n", "\n".join(out))
    return result.strip() + "\n"


# --------------------------------------------------------------------------
# Metadata recognition
# --------------------------------------------------------------------------

# The title/наказ number/date always live in the cover-page approval block
# ("ЗАТВЕРДЖЕНО\nНаказ ... від DD MONTH YYYY р. № NNN/од\n...ПОЛОЖЕННЯ..."),
# right after the word ЗАТВЕРДЖЕНО. Searching only that narrow window (not
# the whole document, and not even just "the first N characters") avoids
# matching an unrelated date/№ cited in body text (e.g. "...затверджено
# наказом МОН від 06.10.2010 №...", a legal citation inside section 1.1).
COVER_TEXT_CHARS = 2500
APPROVAL_SLICE_CHARS = 600
RE_APPROVAL_WORD = re.compile(r"ЗАТВЕРДЖЕНО")


def _approval_slice(text: str) -> str:
    m = RE_APPROVAL_WORD.search(text)
    if m:
        return text[m.start(): m.start() + APPROVAL_SLICE_CHARS]
    return text[:COVER_TEXT_CHARS]

RE_ANY_NUMBER_SIGN = re.compile(r"№\s*([\w./-]+)")
# "№" preceded by these labels refers to a copy/page/protocol/table-row
# count, not the наказ's own registration number (e.g. document headers
# all carry "Екземпляр № 1" on every page, and tables use "№ з/п" as a
# "row no." column header).
NUMBER_SIGN_STOPWORDS = re.compile(
    r"(екземпляр|арк\.?|випуск|протокол|з/п)\s*$", re.IGNORECASE
)

MONTHS_UK = {
    "січня": "01", "лютого": "02", "березня": "03", "квітня": "04",
    "травня": "05", "червня": "06", "липня": "07", "серпня": "08",
    "вересня": "09", "жовтня": "10", "листопада": "11", "грудня": "12",
}
RE_ORDER_DATE_TEXTUAL = re.compile(
    r"(\d{1,2})\s+(" + "|".join(MONTHS_UK) + r")\s+(\d{4})",
    re.IGNORECASE,
)
RE_ORDER_DATE_TEXTUAL_AFTER_VID = re.compile(
    r"від\s+(\d{1,2})\s+(" + "|".join(MONTHS_UK) + r")\s+(\d{4})",
    re.IGNORECASE,
)
RE_ORDER_DATE_NUMERIC = re.compile(r"\b(\d{2})[./](\d{2})[./](\d{4})\b")


def guess_order_number(text: str) -> str | None:
    cover = _approval_slice(text)
    for m in RE_ANY_NUMBER_SIGN.finditer(cover):
        context_before = cover[max(0, m.start() - 25): m.start()]
        value = m.group(1).strip().rstrip(".,")
        if NUMBER_SIGN_STOPWORDS.search(context_before) or not any(ch.isdigit() for ch in value):
            continue
        return value
    return None


def guess_order_date(text: str) -> str | None:
    cover = _approval_slice(text)
    # Prefer a date right after "від" (наказ ... від DD <місяць> YYYY р.),
    # since a document can also contain unrelated dates (e.g. a Вчена рада
    # protocol date, or a citation of some other regulation) elsewhere.
    m = RE_ORDER_DATE_TEXTUAL_AFTER_VID.search(cover)
    if m:
        day, month_name, year = m.groups()
        return f"{year}-{MONTHS_UK[month_name.lower()]}-{int(day):02d}"
    m = RE_ORDER_DATE_TEXTUAL.search(cover)
    if m:
        day, month_name, year = m.groups()
        month = MONTHS_UK[month_name.lower()]
        return f"{year}-{month}-{int(day):02d}"
    m = RE_ORDER_DATE_NUMERIC.search(cover)
    if m:
        day, month, year = m.groups()
        return f"{year}-{month}-{day}"
    return None


# Document-type words that open the title, either alone on their own line
# ("ПОЛОЖЕННЯ", with the rest of the title on following lines) or as the
# first word of an all-caps title line ("КОДЕКС АКАДЕМІЧНОЇ ДОБРОЧЕСНОСТІ").
TITLE_ANCHORS = {
    "ПОЛОЖЕННЯ", "НАКАЗ", "ІНСТРУКЦІЯ", "ПРАВИЛА", "ПОРЯДОК", "РЕГЛАМЕНТ",
    "КОДЕКС", "ПРОЦЕДУРА", "ПОЛІТИКА", "СТАТУТ", "СТРАТЕГІЯ",
    "ПРОГРАМА", "ПЕРЕЛІК", "РЕЄСТР", "ЗВІТ", "КОНЦЕПЦІЯ", "НАСТАНОВА",
    "КАРТА", "ПЛАН-ГРАФІК", "ІНДЕКСАЦІЯ", "РОЗПОРЯДЖЕННЯ",
}
# Lines that signal we've run past the title into unrelated boilerplate.
# "ректор" is word-bounded because plenty of real titles legitimately
# inflect it ("вибори ректора", "почесного ректора") — only the standalone
# "Ректор" signature label should stop collection, not every word built on
# that stem.
TITLE_STOP_LINE = re.compile(
    r"(контрольний примірник|врахований примірник|погоджено|вченою радою|\bректор\b|_{3,})",
    re.IGNORECASE,
)


def _is_mostly_uppercase(line: str) -> bool:
    letters = [ch for ch in line if ch.isalpha()]
    if not letters:
        return False
    return sum(ch.isupper() for ch in letters) / len(letters) > 0.8


def _find_title_anchor_parts(text: str) -> tuple[str, str] | None:
    """If a TITLE_ANCHORS word opens the title — either as the line's first
    word ("ПОЛОЖЕННЯ...") or, when a lead adjective comes first, as its
    second ("ЕКОЛОГІЧНА ПОЛІТИКА...") — return (anchor_word, full_title).
    Used by guess_title() to recover a title from raw PDF text in the
    first place; build_link_spec() does the equivalent split on the final
    frontmatter title string instead, via _split_title_anchor()."""
    lines = [ln.strip() for ln in _approval_slice(text).split("\n")]

    for i, line in enumerate(lines):
        if not line or not _is_mostly_uppercase(line):
            continue
        if line.strip() == "ЗАГАЛЬНІ ПОЛОЖЕННЯ":
            # Fixed idiom for a document's own opening section ("General
            # Provisions"), never a real document title — every single
            # "ПОЛОЖЕННЯ про ..." document in this corpus opens its body
            # with this exact heading, so treating ПОЛОЖЕННЯ-as-2nd-word
            # here would mistake section 1 for the cover title.
            continue
        words = line.split(maxsplit=2)
        anchor_word = next((w for w in words[:2] if w in TITLE_ANCHORS), None)
        if anchor_word:
            full_title = _collect_title(lines, i + 1, line)
            return anchor_word, full_title
    return None


def _looks_like_split_logo_or_code(line: str) -> bool:
    """The institution's logo ("Житомирська" / "політехніка") sometimes
    extracts as its own line, alone or glued to a document-control code
    ("Житомирська П 18.00 - 05 - 2019") — STANDALONE_NOISE_LINES only
    catches the bare-word case. Treat a line as this kind of noise if it's
    just one of those two words, or if one of them appears alongside at
    least two digits (a document code fragment sharing the line). The
    length cap keeps this from misfiring on a genuine title/sentence that
    happens to name the institution and also contain a year ("...у
    Державному університеті «Житомирська політехніка» на 2024-2026 роки")
    — real glued-logo fragments are always short."""
    low = line.lower().strip()
    if low in {"житомирська", "політехніка"}:
        return True
    if len(line) > 60:
        return False
    words = set(re.findall(r"[а-яіїєґ]+", low))
    digit_count = sum(ch.isdigit() for ch in line)
    return bool(words & {"житомирська", "політехніка"}) and digit_count >= 2


def _looks_like_control_code(line: str) -> bool:
    """A document-control code line ("П – 04.00 – 01/04-02 – 2024") is
    almost all digits/punctuation with at most one or two stray letters —
    unlike real title text, which is letter-heavy even when it names a
    year."""
    letters = [ch for ch in line if ch.isalpha()]
    digit_count = sum(ch.isdigit() for ch in line)
    return digit_count >= 4 and len(letters) <= 2


# The ЗАТВЕРДЖЕНО block's "issuing order" line ("Наказ Державного
# університету «Житомирська політехніка»") and the bare quoted institution
# name that often follows it are boilerplate that separates the anchor word
# from the real title text, or (when no anchor word is found at all) sit
# between ЗАТВЕРДЖЕНО and the real, non-uppercase title line further down.
RE_APPROVAL_ISSUER_LINE = re.compile(
    r"^Наказ\s+(Державного\s+ун[іi]верситету|М[іi]н[іi]стерства)", re.IGNORECASE
)
RE_INSTITUTION_NAME_ONLY = re.compile(
    r'^[«"]?\s*Житомирська\s+політехніка\s*[»"]?[.,]?$', re.IGNORECASE
)


def _is_transient_title_noise(line: str) -> bool:
    """Noise that can appear *between* real title-text lines and should be
    skipped without ending the title (a document-control code, an order
    date/number, the issuing-order line) — as opposed to TITLE_STOP_LINE,
    which marks a firm end to the title block. Deliberately excludes
    RE_INSTITUTION_NAME_ONLY: that line is boilerplate before the title
    starts, but a legitimate title routinely *ends* with exactly the
    quoted institution name on its own line — callers that haven't
    started collecting real content yet should check it separately."""
    return (
        bool(RE_DOT_LEADER.search(line))
        or _looks_like_split_logo_or_code(line)
        or _looks_like_control_code(line)
        or bool(RE_ORDER_DATE_TEXTUAL.search(line))
        or bool(RE_ORDER_DATE_NUMERIC.search(line))
        or bool(RE_ANY_NUMBER_SIGN.search(line))
        or bool(RE_APPROVAL_ISSUER_LINE.match(line))
    )


# How many lines past the anchor/start line to keep looking for more title
# text. Generous enough to cover a multi-line "НАКАЗ" subject paragraph
# (which, unlike a "ПОЛОЖЕННЯ" title, can run to 6-8 lines), while the
# blank-line/TITLE_STOP_LINE break below still keeps it from running past
# the title into the document body.
TITLE_CONTINUATION_WINDOW = 16

# A document's own first body section ("1. Загальні положення", "I. ЗАГАЛЬНІ
# ПОЛОЖЕННЯ") starts with a bare number/roman numeral + dot. Some documents
# have no ЗАТВЕРДЖЕНО/Контрольний примірник separator at all between their
# bare anchor word and this heading, so without treating it as a firm stop
# too, title collection runs straight into the body. Ukrainian text types
# roman numerals with the Cyrillic lookalikes І/Х (U+0406/U+0425), not the
# visually identical Latin I/X — both need to match here.
RE_BODY_SECTION_START = re.compile(r"^(?:\d{1,3}|[IVXLCDMІХ]{1,4})\.\s+\S")


def _collect_title(lines: list[str], start_idx: int, first_line: str) -> str:
    """Build a full (possibly multi-line) title starting at first_line,
    extending forward through lines[start_idx:] while skipping transient
    noise, and stopping firmly at a blank line once real content has been
    captured, or at a TITLE_STOP_LINE/RE_BODY_SECTION_START marker (which
    never re-opens — everything past it is signature/approval boilerplate
    or the document body, never title text, however it happens to be
    interleaved)."""
    parts = [first_line]
    content_started = False
    for cont in lines[start_idx: start_idx + TITLE_CONTINUATION_WINDOW]:
        if TITLE_STOP_LINE.search(cont) or RE_BODY_SECTION_START.match(cont):
            break
        if not cont:
            if content_started:
                break
            continue
        if not content_started and RE_INSTITUTION_NAME_ONLY.match(cont):
            continue
        if _is_transient_title_noise(cont):
            continue
        parts.append(cont)
        content_started = True
    return " ".join(parts)[:300]


def guess_title(text: str) -> str:
    lines = [ln.strip() for ln in _approval_slice(text).split("\n")]

    found = _find_title_anchor_parts(text)
    if found:
        return found[1]

    # Fallback: first substantial line that isn't a repeated header, a TOC
    # dot-leader entry, a bare document-code fragment, or ЗАТВЕРДЖЕНО-block
    # boilerplate. Some documents in this corpus never carry an anchor word
    # ("ПОЛОЖЕННЯ", "НАКАЗ", ...) at the start of their title line — so this
    # is genuinely the best available signal for them.  RUNNING_HEADER_PATTERNS
    # is written for strip_boilerplate, where case is the ALL-CAPS-vs-title-case
    # signal that distinguishes boilerplate from real content; here we match
    # it case-insensitively, since some documents render "Міністерство освіти
    # і науки України" in title case instead of the more common ALL-CAPS, and
    # that boilerplate is just as unhelpful as a title either way.
    for i, line in enumerate(lines):
        if (
            len(line) >= 8
            and not _is_mostly_uppercase(line)
            and any(ch.isalpha() for ch in line)
            and not RE_DOT_LEADER.search(line)
            and not _looks_like_split_logo_or_code(line)
            and not _is_transient_title_noise(line)
            and not RE_INSTITUTION_NAME_ONLY.match(line)
            and not TITLE_STOP_LINE.search(line)
            and not RE_BODY_SECTION_START.match(line)
            and not any(re.search(p.pattern, line, re.IGNORECASE) for p in RUNNING_HEADER_PATTERNS)
        ):
            return _collect_title(lines, i + 1, line[:200])
    return ""


# --------------------------------------------------------------------------
# Cross-document linking
# --------------------------------------------------------------------------

# Ukrainian declines "Положення про X" into "Положенням про X" / "Положенні
# про X" / etc. depending on grammatical role, but the noun stem stays
# fixed, so "<stem>\w*" matches every case. "Положення" itself barely
# declines (only dative/instrumental/locative change), but the others do.
ANCHOR_STEMS = {
    "ПОЛОЖЕННЯ": "Положен",
    "НАКАЗ": "Наказ",
    "ІНСТРУКЦІЯ": "Інструкці",
    "ПРАВИЛА": "Правил",
    "ПОРЯДОК": "Поряд",
    "РЕГЛАМЕНТ": "Регламент",
    "КОДЕКС": "Кодекс",
    "ПРОЦЕДУРА": "Процедур",
    "ПОЛІТИКА": "Політик",
    "СТАТУТ": "Статут",
    "СТРАТЕГІЯ": "Стратег",
    "ПРОГРАМА": "Програм",
    "ПЕРЕЛІК": "Перелі",
    "РЕЄСТР": "Реєстр",
    "ЗВІТ": "Звіт",
    "КОНЦЕПЦІЯ": "Концепці",
    "НАСТАНОВА": "Настанов",
    "КАРТА": "Карт",
    "ПЛАН-ГРАФІК": "План-граф",
    "ІНДЕКСАЦІЯ": "Індексаці",
    "РОЗПОРЯДЖЕННЯ": "Розпоряджен",
}

# The "Державного університету «Житомирська політехніка»" tail is common to
# almost every title and appears in body references with varying grammar
# ("у Державному університеті...", "Державного університету..."), or is
# dropped entirely when the reference is abbreviated. Strip it off the end
# of a title's core phrase, and match any inflected form of it as optional
# in body text, rather than requiring the title's exact wording.
RE_UNIV_SUFFIX = re.compile(
    r"\s*(?:у\s+)?Державн\w+\s+університет\w*\s+«Житомирська\s+політехніка»\.?$",
    re.IGNORECASE,
)
RE_UNIV_SUFFIX_MATCH = r"(?:\s+\S+)?\s+Державн\w+\s+університет\w*\s+«Житомирська\s+політехніка»"

# A core phrase shorter than this is too generic to link safely (risk of
# matching unrelated text).
MIN_LINK_CORE_WORDS = 2


@dataclass
class LinkSpec:
    pattern: re.Pattern
    url: str


def _split_title_anchor(title: str) -> tuple[str, str, str] | None:
    """Find a TITLE_ANCHORS word opening a title *string* — as it finally
    ended up in frontmatter, title-case (a hand-simplified title) or
    ALL-CAPS (machine-guessed) — rather than re-deriving one from the raw
    PDF text the way _find_title_anchor_parts does. This is what makes a
    document linkable as a *target*: even one whose title had to be
    hand-fixed (no anchor recoverable from its own raw cover-page text at
    all) can still be linked to from other documents' bodies, since by
    now every document has *some* title. Returns (anchor_word, lead,
    rest): lead is any adjective before a second-word anchor ("Дорожня"
    in "Дорожня карта...", "ЕКОЛОГІЧНА" in "ЕКОЛОГІЧНА ПОЛІТИКА...") —
    empty when the anchor opens the title — and rest is everything after
    the anchor."""
    words = title.strip().split()
    for i in range(min(2, len(words))):
        if words[i].upper() in TITLE_ANCHORS:
            return words[i].upper(), " ".join(words[:i]), " ".join(words[i + 1:])
    return None


def build_link_spec(pdf_stem: str, title: str) -> LinkSpec | None:
    """Build a regex that finds mentions of this document (by its own
    title) inside OTHER documents' bodies, so they can become links."""
    found = _split_title_anchor(title)
    if not found:
        return None
    anchor_word, lead, rest = found
    stem = ANCHOR_STEMS.get(anchor_word)
    if not stem:
        return None

    core = RE_UNIV_SUFFIX.sub("", rest).strip()
    lead_words = lead.split()
    core_words = core.split()
    if len(lead_words) + len(core_words) < MIN_LINK_CORE_WORDS:
        return None

    # A mention inside another document's body is ordinary sentence-case
    # prose regardless of how *this* document's own title happens to be
    # cased (ALL-CAPS straight off an anchor-detected cover, or title-case
    # after a hand fix) — match case-insensitively so either source still
    # finds it.
    lead_pattern = r"\s+".join(re.escape(w) for w in lead_words)
    core_pattern = r"\s+".join(re.escape(w) for w in core_words)
    anchor_part = rf"{lead_pattern}\s+{stem}\w*" if lead_pattern else rf"{stem}\w*"
    body_part = rf"{anchor_part}\s+{core_pattern}" if core_pattern else anchor_part
    pattern = re.compile(rf"{body_part}(?:{RE_UNIV_SUFFIX_MATCH})?", re.IGNORECASE)
    url = SITE_BASEURL + "/wiki-pages/" + quote(f"{pdf_stem}.html", safe="/")
    return LinkSpec(pattern=pattern, url=url)


def link_cross_references(body: str, own_stem: str, registry: dict[str, LinkSpec]) -> str:
    """Turn the first mention of each other recognized document's title
    into a Markdown link. Non-overlapping, self-links excluded."""
    matches: list[tuple[int, int, str]] = []
    for target_stem, spec in registry.items():
        if target_stem == own_stem:
            continue
        m = spec.pattern.search(body)
        if m:
            matches.append((m.start(), m.end(), spec.url))

    if not matches:
        return body

    matches.sort()
    accepted: list[tuple[int, int, str]] = []
    last_end = -1
    for start, end, url in matches:
        if start < last_end:
            continue
        accepted.append((start, end, url))
        last_end = end

    out = []
    cursor = 0
    for start, end, url in accepted:
        out.append(body[cursor:start])
        out.append(f"[{body[start:end]}]({url})")
        cursor = end
    out.append(body[cursor:])
    return "".join(out)


# --------------------------------------------------------------------------
# Markdown assembly
# --------------------------------------------------------------------------

def read_existing_frontmatter(out_path: Path) -> dict:
    """Re-running the script must not clobber fields a human or a later
    pipeline step (e.g. scripts/categorize.py) already set on a previous
    .md output — status, category, subgroup, and (only when explicitly
    hand-verified via title_locked: true) title. Only brand-new files get
    status 'невідомо' and no category."""
    if out_path.exists():
        try:
            raw = out_path.read_text(encoding="utf-8")
            if raw.startswith("---\n"):
                fm_text = raw.split("---\n", 2)[1]
                existing = yaml.safe_load(fm_text) or {}
                return {
                    "status": existing.get("status") or "невідомо",
                    "category": existing.get("category"),
                    "subgroup": existing.get("subgroup"),
                    "title": existing.get("title") if existing.get("title_locked") else None,
                }
        except Exception:
            pass
    return {"status": "невідомо", "category": None, "subgroup": None, "title": None}


@dataclass
class DocRecord:
    source_path: Path
    out_path: Path
    result: ExtractionResult
    existing_status: str
    existing_category: str | None
    existing_subgroup: str | None
    locked_title: str | None
    title: str
    order_number: str
    order_date: str
    body: str
    link_spec: LinkSpec | None


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------

def analyze_document(source_path: Path, stats: RunStats) -> DocRecord | None:
    """Phase 1: extract + guess metadata for one source document (PDF or
    DOCX). Cross-document links aren't inserted yet — that needs every
    document's title first, which only exists once phase 1 has run for all
    of them."""
    stats.total += 1
    try:
        result = extract_document(source_path)
    except Exception as exc:
        log.error("FAILED %s: %s", source_path.name, exc)
        stats.failed.append((source_path.name, str(exc)))
        return None

    if not result.text.strip():
        msg = "no text extracted"
        log.error("FAILED %s: %s", source_path.name, msg)
        stats.failed.append((source_path.name, msg))
        return None

    out_path = OUTPUT_DIR / f"{source_path.stem}.md"
    body_text = "\n\n".join(drop_cover_page(result.pages_text))
    body = split_into_blocks(strip_trailing_administrivia(strip_boilerplate(body_text)))
    existing = read_existing_frontmatter(out_path)
    title = existing["title"] or guess_title(result.text)

    return DocRecord(
        source_path=source_path,
        out_path=out_path,
        result=result,
        existing_status=existing["status"],
        existing_category=existing["category"],
        existing_subgroup=existing["subgroup"],
        locked_title=existing["title"],
        title=title,
        order_number=guess_order_number(result.text) or "",
        order_date=guess_order_date(result.text) or "",
        body=body,
        link_spec=build_link_spec(source_path.stem, title),
    )


def write_record(record: DocRecord, registry: dict[str, LinkSpec]) -> int:
    """Phase 2: insert cross-document links (now that the full registry
    exists) and write the final .md file. Returns the number of links
    inserted, for the run summary."""
    linked_body = link_cross_references(record.body, record.source_path.stem, registry)
    links_inserted = linked_body.count(f"]({SITE_BASEURL}/wiki-pages/")

    frontmatter = {
        "title": record.title,
        "status": record.existing_status,
        "order_number": record.order_number,
        "order_date": record.order_date,
        "source_pdf": record.source_path.name,
    }
    if record.locked_title:
        frontmatter["title_locked"] = True
    if record.existing_category:
        frontmatter["category"] = record.existing_category
    if record.existing_subgroup:
        frontmatter["subgroup"] = record.existing_subgroup
    fm_yaml = yaml.dump(frontmatter, allow_unicode=True, sort_keys=False, default_flow_style=False)
    md = f"---\n{fm_yaml}---\n\n{linked_body}"
    record.out_path.write_text(md, encoding="utf-8")
    return links_inserted


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    source_files = sorted(SOURCE_DIR.glob("*.pdf")) + sorted(SOURCE_DIR.glob("*.docx"))
    if not source_files:
        log.warning("No PDF/DOCX files found in %s", SOURCE_DIR)
        return 0

    stats = RunStats()
    records = [r for r in (analyze_document(p, stats) for p in source_files) if r is not None]

    registry = {r.source_path.stem: r.link_spec for r in records if r.link_spec}
    log.info("Cross-reference registry: %d/%d documents recognized a linkable title", len(registry), len(records))

    total_links = 0
    for record in records:
        total_links += write_record(record, registry)

        stats.succeeded += 1
        if record.result.used_ocr:
            stats.used_ocr += 1
        log.info(
            "OK %s -> %s (%d pages, %s)",
            record.source_path.name, record.out_path.name, record.result.pages,
            "OCR" if record.result.used_ocr else "text layer",
        )

    print("\n=== Підсумок ===")
    print(f"Усього документів:     {stats.total}")
    print(f"Оброблено успішно:     {stats.succeeded}")
    print(f"Знадобився OCR:        {stats.used_ocr}")
    print(f"Розпізнано для лінків: {len(registry)}")
    print(f"Проставлено посилань:  {total_links}")
    print(f"Помилки:               {len(stats.failed)}")
    for name, err in stats.failed:
        print(f"  - {name}: {err}")

    return 1 if stats.failed else 0


if __name__ == "__main__":
    sys.exit(main())
